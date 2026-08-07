import pdfplumber
from docx import Document
from pathlib import Path


def extract_pdf_text(file_path):
    """
    Extract text from a PDF resume.
    """
    print(f"\nReading PDF: {file_path}")

    text = ""

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

    return text


def extract_docx_text(file_path):
    """
    Extract text from a DOCX resume.
    """
    print(f"\nReading DOCX: {file_path}")

    doc = Document(file_path)

    print(f"Paragraphs found: {len(doc.paragraphs)}")

    text = ""

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text.strip() + "\n"

    return text


def extract_resume_text(file_path):
    """
    Extract text from a resume (PDF or DOCX).
    """

    file_path = Path(file_path)

    print(f"\nSelected file: {file_path}")
    print(f"Extension: {file_path.suffix}")

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if file_path.suffix.lower() == ".pdf":
        return extract_pdf_text(file_path)

    elif file_path.suffix.lower() == ".docx":
        return extract_docx_text(file_path)

    else:
        raise ValueError(f"Unsupported file format: {file_path.suffix}")